#### Back EMF Tester
#### Doug Gammill
#### Linear Labs Inc
#### April 19, 2022

import configTemplate
import SDS1104XE # handles Oscopes
import arduino # handles Relays
from tkinter import * # handles GUI
from threading import * # handles threading so GUI doesnt freeze during main functions
import tkinter as tk # handles GUI
from PIL import Image, ImageTk # handles GUI Image
from tkinter import filedialog # handles opening Results and Config file paths
import os  # for working with opening files with applications like excel and notepad
import configparser # handles config files
from pathlib import Path # handles (if Path(configFilePath).is_file():)
import pathlib
import pandas as pd # handles data frames
from pandas import ExcelWriter # used with saving spread sheets with sheet numbers
import openpyxl
#from time import sleep # handles delay
import time # used for adding time stamps
import docx # for editing documents with images (pip install python-docx)
from docx.shared import Inches # docx conversion to inches

passFail = 0 # if error, use this index to skip back to main
saveIndex=0 # 0 = end 
comOK=0 # 0 = usb not connected
testIndex=0 # determines if testing or not
barcode=''
configFolder = configTemplate.configFolder # shared config file path for motors
folder = configTemplate.resultsSavePath # dataLog savePath
screenshotsSavePath = configTemplate.screenshotsSavePath # screenshots savePath
skipButton = configTemplate.skipButton

def readConfig():
    global partNum,passFail,status,configFolder,skipFail,noHalls,stator,speedMin,speedMax,aVrmsMin,aVrmsMax,bVrmsMin,bVrmsMax,cVrmsMin,cVrmsMax,abPhaseMin,abPhaseMax,acPhaseMin,acPhaseMax,bcPhaseMin,bcPhaseMax,aHallPhaseMin,aHallPhaseMax,bHallPhaseMin,bHallPhaseMax,cHallPhaseMin,cHallPhaseMax,aHallDutyMin,aHallDutyMax,bHallDutyMin,bHallDutyMax,cHallDutyMin,cHallDutyMax,aHallVoltMin,aHallVoltMax,bHallVoltMin,bHallVoltMax,cHallVoltMin,cHallVoltMax,abHallMin,abHallMax,acHallMin,acHallMax,bcHallMin,bcHallMax,triggerLevel1,triggerLevel2,triggerCoupling1,triggerCoupling2,triggerChannel1,triggerChannel2,triggerMode1,triggerMode2,triggerSelect1,triggerSelect2,channelOffset1,channelOffset2,voltageDivision1,voltageDivision2,channelUnits1,channelUnits2,timeDivision1,timeDivision2,memDepth,x10Probe
    # get partNum params
    configParser = configparser.RawConfigParser()   
    configFilePath = configFolder+'/'+partNum+'.txt'
    print(configFilePath)
    
    if Path(configFilePath).is_file():
        configParser.read(configFilePath)
        
        skipFail = configParser.get(partNum, 'skipFail')
        skipFail=int(skipFail)
        
        noHalls = configParser.get(partNum, 'noHalls')
        noHalls=int(noHalls)
        
        speedMin = configParser.get(partNum, 'speedMin') # hallA minimum speed
        speedMin=float(speedMin)
        
        speedMax = configParser.get(partNum, 'speedMax') # hallA max speed
        speedMax=float(speedMax)
        
        aVrmsMin = configParser.get(partNum, 'aVrmsMin') # phaseA Vrms minimum
        aVrmsMin=float(aVrmsMin)
        
        aVrmsMax = configParser.get(partNum, 'aVrmsMax') # phaseA Vrms maximum 
        aVrmsMax=float(aVrmsMax)
        
        bVrmsMin = configParser.get(partNum, 'bVrmsMin') # phaseB Vrms minimum
        bVrmsMin=float(bVrmsMin)
        
        bVrmsMax = configParser.get(partNum, 'bVrmsMax') # phaseB Vrms maximum
        bVrmsMax=float(bVrmsMax)
        
        cVrmsMin = configParser.get(partNum, 'cVrmsMin') # phaseC Vrms minimum
        cVrmsMin=float(cVrmsMin)
        
        cVrmsMax = configParser.get(partNum, 'cVrmsMax') # phaseC Vrms maximum
        cVrmsMax=float(cVrmsMax)
        
        abPhaseMin = configParser.get(partNum, 'abPhaseMin') # AB phase to phase offset in degrees min
        abPhaseMin=float(abPhaseMin)
        
        abPhaseMax = configParser.get(partNum, 'abPhaseMax') # AB phase to phase offset in degrees max
        abPhaseMax=float(abPhaseMax)
        
        acPhaseMin = configParser.get(partNum, 'acPhaseMin') # AC phase to phase offset in degrees min
        acPhaseMin=float(acPhaseMin)
        
        acPhaseMax = configParser.get(partNum, 'acPhaseMax') # AC phase to phase offset in degrees max
        acPhaseMax=float(acPhaseMax)
        
        bcPhaseMin = configParser.get(partNum, 'bcPhaseMin') # BC phase to phase offset in degrees min
        bcPhaseMin=float(bcPhaseMin)
        
        bcPhaseMax = configParser.get(partNum, 'bcPhaseMax') # BC phase to phase offset in degrees max
        bcPhaseMax=float(bcPhaseMax)
        
        aHallPhaseMin = configParser.get(partNum, 'aHallPhaseMin') # hallA to phaseA alignment in degrees min
        aHallPhaseMin=float(aHallPhaseMin)
        
        aHallPhaseMax = configParser.get(partNum, 'aHallPhaseMax') # hallA to phaseA alignment in degrees max
        aHallPhaseMax=float(aHallPhaseMax)

        bHallPhaseMin = configParser.get(partNum, 'bHallPhaseMin') # hallB to phaseB alignment in degrees min
        bHallPhaseMin=float(bHallPhaseMin)
        
        bHallPhaseMax = configParser.get(partNum, 'bHallPhaseMax') # hallB to phaseB alignment in degrees max
        bHallPhaseMax=float(bHallPhaseMax)

        cHallPhaseMin = configParser.get(partNum, 'cHallPhaseMin') # hallC to phaseC alignment in degrees min
        cHallPhaseMin=float(cHallPhaseMin)
        
        cHallPhaseMax = configParser.get(partNum, 'cHallPhaseMax') # hallC to phaseC alignment in degrees max
        cHallPhaseMax=float(cHallPhaseMax)
        
        aHallDutyMin = configParser.get(partNum, 'aHallDutyMin') # hallA duty cycle min
        aHallDutyMin=float(aHallDutyMin)
        
        aHallDutyMax = configParser.get(partNum, 'aHallDutyMax') # hallA duty cycle max
        aHallDutyMax=float(aHallDutyMax)
        
        bHallDutyMin = configParser.get(partNum, 'bHallDutyMin') # hallB duty cycle min
        bHallDutyMin=float(bHallDutyMin)
        
        bHallDutyMax = configParser.get(partNum, 'bHallDutyMax') # hallB duty cycle max
        bHallDutyMax=float(bHallDutyMax)
        
        cHallDutyMin = configParser.get(partNum, 'cHallDutyMin') # hallC duty cycle min
        cHallDutyMin=float(cHallDutyMin)
        
        cHallDutyMax = configParser.get(partNum, 'cHallDutyMax') # hallC duty cycle max
        cHallDutyMax=float(cHallDutyMax)
        
        aHallVoltMin = configParser.get(partNum, 'aHallVoltMin') # hallA Volt min
        aHallVoltMin=float(aHallVoltMin)
        
        aHallVoltMax = configParser.get(partNum, 'aHallVoltMax') # hallA Volt max
        aHallVoltMax=float(aHallVoltMax)
        
        bHallVoltMin = configParser.get(partNum, 'bHallVoltMin') # hallB Volt min
        bHallVoltMin=float(bHallVoltMin)
        
        bHallVoltMax = configParser.get(partNum, 'bHallVoltMax') # hallB Volt max
        bHallVoltMax=float(bHallVoltMax)
        
        cHallVoltMin = configParser.get(partNum, 'cHallVoltMin') # hallC Volt min
        cHallVoltMin=float(cHallVoltMin)
        
        cHallVoltMax = configParser.get(partNum, 'cHallVoltMax') # hallC Volt max
        cHallVoltMax=float(cHallVoltMax)
        
        abHallMin = configParser.get(partNum, 'abHallMin') # AB Hall to Hall offset in degrees min
        abHallMin=float(abHallMin)
        
        abHallMax = configParser.get(partNum, 'abHallMax') # AB Hall to Hall offset in degrees max
        abHallMax=float(abHallMax)
        
        acHallMin = configParser.get(partNum, 'acHallMin') # AC Hall to Hall offset in degrees min
        acHallMin=float(acHallMin)
        
        acHallMax = configParser.get(partNum, 'acHallMax') # AC Hall to Hall offset in degrees max
        acHallMax=float(acHallMax)
        
        bcHallMin = configParser.get(partNum, 'bcHallMin') # BC Hall to Hall offset in degrees min
        bcHallMin=float(bcHallMin)
        
        bcHallMax = configParser.get(partNum, 'bcHallMax') # BC Hall to Hall offset in degrees max
        bcHallMax=float(bcHallMax)

        triggerChannel1 = configParser.get(partNum, 'triggerChannel1')
        triggerChannel2 = configParser.get(partNum, 'triggerChannel2')
        
        triggerLevel1 = configParser.get(partNum, 'triggerLevel1')
        triggerLevel2 = configParser.get(partNum, 'triggerLevel2')
                
        triggerCoupling1 = configParser.get(partNum, 'triggerCoupling1')
        triggerCoupling2 = configParser.get(partNum, 'triggerCoupling2')
                
        triggerMode1 = configParser.get(partNum, 'triggerMode1')
        triggerMode2 = configParser.get(partNum, 'triggerMode2')
                
        triggerSelect1 = configParser.get(partNum, 'triggerSelect1')
        triggerSelect2 = configParser.get(partNum, 'triggerSelect2')
                
        channelOffset1 = configParser.get(partNum, 'channelOffset1')
        channelOffset2 = configParser.get(partNum, 'channelOffset2')
                
        voltageDivision1 = configParser.get(partNum, 'voltageDivision1')
        voltageDivision2  = configParser.get(partNum, 'voltageDivision2')
                
        channelUnits1 = configParser.get(partNum, 'channelUnits1')
        channelUnits2 = configParser.get(partNum, 'channelUnits2')
                
        timeDivision1 = configParser.get(partNum, 'timeDivision1') 
        timeDivision2  = configParser.get(partNum, 'timeDivision2')
                
        memDepth  = configParser.get(partNum, 'memDepth')
                
        x10Probe = configParser.get(partNum, 'x10Probe')

        stator = configParser.get(partNum, 'stator')
        
        
    else:
        print('ERROR: No Config File Found!')
        status = 'ERROR: No Config File Found!'
        text.config(text=status, bg='red', fg='black')
        passFail=0
        return passFail

def autoSetup():
    global Scope1,Scope2,partNum
    SDS1104XE.autoSetup(Scope1)
    if noHalls==0:
        SDS1104XE.autoSetup(Scope2)
    i=15
    while i>0:
        text.config(text='Please wait while Scope(s) are setup '+str(i), bg='Yellow', fg='black')
        i=i-1

    text.config(fg="black",bg='Yellow',width=25,font=12, text='Testing: '+partNum)

def setupScopes():
    global Scope1,Scope2,noHalls,triggerChannel1,triggerChannel2,triggerLevel1,triggerLevel2,triggerCoupling1,triggerCoupling2,triggerMode1,triggerMode2,triggerSelect1,triggerSelect2,channelOffset1,channelOffset2,voltageDivision1,voltageDivision2,channelUnits1,channelUnits2,timeDivision1,timeDivision2,memDepth,x10Probe

    text.config(fg="black",bg='orange',width=25,font='12', text='Setting Up Scopes')
    SDS1104XE.setupScope1(Scope1,triggerChannel1,triggerLevel1,triggerCoupling1,triggerMode1,triggerSelect1,channelOffset1,voltageDivision1,channelUnits1,timeDivision1,memDepth,x10Probe)
    if noHalls == 0:
        SDS1104XE.setupScope2(Scope2,triggerChannel2,triggerLevel2,triggerCoupling2,triggerMode2,triggerSelect2,channelOffset2,voltageDivision2,channelUnits2,timeDivision2,memDepth,x10Probe)
        SDS1104XE.turnOffChannel1(Scope1)
        
def getValuesScope1(): # Phase
    global passFail,noHalls,voltageDivision2,channelUnits2,x10Probe,failureMessage,failureMessage2,failureMessage3,failureMessage4,failureMessage5,failureMessage6,failureMessage7,failureMessage17,failureMessage18,img2,img3,img4,img5,img6,img7,img8,img18,img19,Scope1,rmsA,rmsB,rmsC,Ph12,Ph13,Ph23,BemfHallPh,BemfHallPhb,BemfHallPhc,aVrmsMin,aVrmsMax,bVrmsMin,bVrmsMax,cVrmsMin,cVrmsMax,abPhaseMin,abPhaseMax,acPhaseMin,acPhaseMax,bcPhaseMin,bcPhaseMax,aHallPhaseMin,aHallPhaseMax,bHallPhaseMin,bHallPhaseMax,cHallPhaseMin,cHallPhaseMax 
    Scope=Scope1
    # Read Hall frequcy and duty cycle
    BackEMF1 = SDS1104XE.measureVoltage(Scope, 'C2')
    BackEMF2 = SDS1104XE.measureVoltage(Scope, 'C3')
    BackEMF3 = SDS1104XE.measureVoltage(Scope, 'C4')

    rmsA=BackEMF1[0]
    rmsB=BackEMF2[0]
    rmsC=BackEMF3[0]
    
    Ph12 = SDS1104XE.measurePhase(Scope, 'C2', 'C3')
    Ph13 = SDS1104XE.measurePhase(Scope, 'C2', 'C4')
    Ph23 = SDS1104XE.measurePhase(Scope, 'C3', 'C4')
    
    if os.path.exists('MotorPhases.png'):
            os.remove('MotorPhases.png')
    SDS1104XE.saveScreen(Scope1,'MotorPhases.png')
        
    if rmsA >=aVrmsMin and rmsA<=aVrmsMax:
        img2=Label(gui, borderwidth=0, image=renderCheck)
        img2.grid(row=4, column=2)
    else:
        img2=Label(gui, borderwidth=0, image=renderX)
        img2.grid(row=4, column=2)
        failureMessage='Failure: Phase A Voltage, '+str(rmsA)+'Vrms'
        text.config(text=failureMessage, bg='red', fg='black')
        passFail=0
        
    if rmsB >=bVrmsMin and rmsB<=bVrmsMax:
        img3=Label(gui, borderwidth=0, image=renderCheck)
        img3.grid(row=5, column=2)
    else:
        img3=Label(gui, borderwidth=0, image=renderX)
        img3.grid(row=5, column=2)
        failureMessage2='Failure: Phase B Voltage, '+str(rmsB)+'Vrms'
        text.config(text=failureMessage2, bg='red', fg='black')
        passFail=0

    if rmsC >=cVrmsMin and rmsC<=cVrmsMax:
        img4=Label(gui, borderwidth=0, image=renderCheck)
        img4.grid(row=6, column=2)
    else:
        img4=Label(gui, borderwidth=0, image=renderX)
        img4.grid(row=6, column=2)
        failureMessage3='Failure: Phase C Voltage, '+str(rmsC)+'Vrms'
        text.config(text=failureMessage3, bg='red', fg='black')
        passFail=0

    if Ph12 >=abPhaseMin and Ph12<=abPhaseMax:
        img5=Label(gui, borderwidth=0, image=renderCheck)
        img5.grid(row=7, column=2)
    else:
        img5=Label(gui, borderwidth=0, image=renderX)
        img5.grid(row=7, column=2)
        failureMessage4='Failure: Phase AB Alignment, '+str(Ph12)+'deg'
        text.config(text=failureMessage4, bg='red', fg='black')
        passFail=0
        
    if Ph13 >=acPhaseMin and Ph13<=acPhaseMax:
        img6=Label(gui, borderwidth=0, image=renderCheck)
        img6.grid(row=8, column=2)
    else:
        img6=Label(gui, borderwidth=0, image=renderX)
        img6.grid(row=8, column=2)
        failureMessage5='Failure: Phase AC Alignment, '+str(Ph13)+'deg'
        text.config(text=failureMessage5, bg='red', fg='black')
        passFail=0

    if Ph23 >=bcPhaseMin and Ph23<=bcPhaseMax:
        img7=Label(gui, borderwidth=0, image=renderCheck)
        img7.grid(row=9, column=2)
    else:
        img7=Label(gui, borderwidth=0, image=renderX)
        img7.grid(row=9, column=2)
        failureMessage6='Failure: Phase BC Alignment, '+str(Ph23)+'deg'
        text.config(text=failureMessage6, bg='red', fg='black')
        passFail=0
    print('PhaseA RMS : ' + str(BackEMF1[0]) + ' Vrms')
    print('PhaseB RMS : ' + str(BackEMF2[0]) + ' Vrms')
    print('PhaseC RMS : ' + str(BackEMF3[0]) + ' Vrms')
    print('Phase A-B '+ str(Ph12) +' Deg')
    print('Phase A-C '+ str(Ph13) +' Deg')
    print('Phase B-C '+ str(Ph23) +' Deg') 

    #### Hall Phase to Motor Phase Alignment
    if noHalls == 0:
        #### A
        arduino.command('rela:1,0,0,0')
        SDS1104XE.setupScope1Channel1(Scope1,channelUnits2,x10Probe)
        SDS1104XE.turnOffChannel3(Scope)
        SDS1104XE.turnOffChannel4(Scope)
        BemfHallPh = SDS1104XE.measurePhase(Scope, 'C1', 'C2')
        if BemfHallPh >=aHallPhaseMin and BemfHallPh<=aHallPhaseMax:
            img8=Label(gui, borderwidth=0, image=renderCheck)
            img8.grid(row=10, column=2)
        else:
            img8=Label(gui, borderwidth=0, image=renderX)
            img8.grid(row=10, column=2)
            failureMessage7='Failure: PhaseA HallA Alignment, '+str(BemfHallPh)+'deg'
            text.config(text=failureMessage7, bg='red', fg='black')
            passFail=0
        print('HallA to PhaseA sensor alignment ' + str(BemfHallPh) +' Deg')
        if os.path.exists('HallPhaseA_To_MotorPhaseA_Alignment.png'):
            os.remove('HallPhaseA_To_MotorPhaseA_Alignment.png')
        SDS1104XE.saveScreen(Scope1,'HallPhaseA_To_MotorPhaseA_Alignment.png') 
        arduino.command('rela:0,0,0,0')

        #### B
        arduino.command('rela:0,1,0,0')
        SDS1104XE.turnOffChannel2(Scope)
        SDS1104XE.turnOnChannel3(Scope)
        BemfHallPhb = SDS1104XE.measurePhase(Scope, 'C1', 'C3')
        if BemfHallPhb >=bHallPhaseMin and BemfHallPhb<=bHallPhaseMax:
            img18=Label(gui, borderwidth=0, image=renderCheck)
            img18.grid(row=11, column=2)
        else:
            img18=Label(gui, borderwidth=0, image=renderX)
            img18.grid(row=11, column=2)
            failureMessage17='Failure: PhaseB HallB Alignment, '+str(BemfHallPhb)+'deg'
            text.config(text=failureMessage17, bg='red', fg='black')
            passFail=0
        print('HallB to PhaseB sensor alignment ' + str(BemfHallPhb) +' Deg')
        if os.path.exists('HallPhaseB_To_MotorPhaseB_Alignment.png'):
            os.remove('HallPhaseB_To_MotorPhaseB_Alignment.png')
        SDS1104XE.saveScreen(Scope1,'HallPhaseB_To_MotorPhaseB_Alignment.png') 
        arduino.command('rela:0,0,0,0')

        #### C
        arduino.command('rela:0,0,1,0')
        SDS1104XE.turnOffChannel3(Scope)
        SDS1104XE.turnOnChannel4(Scope)
        BemfHallPhc = SDS1104XE.measurePhase(Scope, 'C1', 'C4')
        if BemfHallPhc >=cHallPhaseMin and BemfHallPhc<=cHallPhaseMax:
            img19=Label(gui, borderwidth=0, image=renderCheck)
            img19.grid(row=12, column=2)
        else:
            img19=Label(gui, borderwidth=0, image=renderX)
            img19.grid(row=12, column=2)
            failureMessage18='Failure: PhaseC HallC Alignment, '+str(BemfHallPhc)+'deg'
            text.config(text=failureMessage18, bg='red', fg='black')
            passFail=0
        print('HallC to PhaseC sensor alignment ' + str(BemfHallPhc) +' Deg')
        if os.path.exists('HallPhaseC_To_MotorPhaseC_Alignment.png'):
            os.remove('HallPhaseC_To_MotorPhaseC_Alignment.png')
        SDS1104XE.saveScreen(Scope1,'HallPhaseC_To_MotorPhaseC_Alignment.png') 
        arduino.command('rela:0,0,0,0')
        SDS1104XE.turnOffChannel1(Scope)
    
def getValuesScope2(): # Hall
    global passFail,failureMessage8,failureMessage9,failureMessage10,failureMessage11,failureMessage12,failureMessage13,failureMessage14,failureMessage15,failureMessage16,img9,img10,img11,img12,img13,img14,img15,img16,img17,Scope2,DCA,DCB,DCC,voltA,voltB,voltC,Ha12,Ha13,Ha23,aHallDutyMin,aHallDutyMax,bHallDutyMin,bHallDutyMax,cHallDutyMin,cHallDutyMax,aHallVoltMin,aHallVoltMax,bHallVoltMin,bHallVoltMax,cHallVoltMin,cHallVoltMax,abHallMin,abHallMax,acHallMin,acHallMax,bcHallMin,bcHallMax
    Scope=Scope2

    DCA = SDS1104XE.measureDutycycle(Scope, 'C1')#Duty Cycles
    DCB = SDS1104XE.measureDutycycle(Scope, 'C2')
    DCC = SDS1104XE.measureDutycycle(Scope, 'C3')
    
    aHallVolt = SDS1104XE.measureVoltage(Scope, 'C1')#voltages
    bHallVolt = SDS1104XE.measureVoltage(Scope, 'C2')
    cHallVolt = SDS1104XE.measureVoltage(Scope, 'C3')

    voltA=aHallVolt[1]
    voltB=bHallVolt[1]
    voltC=cHallVolt[1]

    Ha12 = SDS1104XE.measurePhase(Scope, 'C1', 'C2')# degrees
    Ha13 = SDS1104XE.measurePhase(Scope, 'C1', 'C3')
    Ha23 = SDS1104XE.measurePhase(Scope, 'C2', 'C3')

    if os.path.exists('HallPhases.png'):
            os.remove('HallPhases.png')
    SDS1104XE.saveScreen(Scope,'HallPhases.png')
    
    if DCA >=aHallDutyMin and DCA<=aHallDutyMax:
        img9=Label(gui, borderwidth=0, image=renderCheck)
        img9.grid(row=4, column=5)
    else:
        img9=Label(gui, borderwidth=0, image=renderX)
        img9.grid(row=4, column=5)
        failureMessage8='Failure: Hall A Duty Cycle, '+str(DCA)+'%'
        text.config(text=failureMessage8, bg='red', fg='black')
        passFail=0
        
    if DCB >=bHallDutyMin and DCB<=bHallDutyMax:
        img10=Label(gui, borderwidth=0, image=renderCheck)
        img10.grid(row=5, column=5)
    else:
        img10=Label(gui, borderwidth=0, image=renderX)
        img10.grid(row=5, column=5)
        failureMessage9='Failure: Hall B Duty Cycle, '+str(DCB)+'%'
        text.config(text=failureMessage9, bg='red', fg='black')
        passFail=0

    if DCC >=cHallDutyMin and DCC<=cHallDutyMax:
        img11=Label(gui, borderwidth=0, image=renderCheck)
        img11.grid(row=6, column=5)
    else:
        img11=Label(gui, borderwidth=0, image=renderX)
        img11.grid(row=6, column=5)
        failureMessage10='Failure: Hall C Duty Cycle, '+str(DCC)+'%'
        text.config(text=failureMessage10, bg='red', fg='black')
        passFail=0

    if voltA >=aHallVoltMin and voltA<=aHallVoltMax:
        img12=Label(gui, borderwidth=0, image=renderCheck)
        img12.grid(row=7, column=5)
    else:
        img12=Label(gui, borderwidth=0, image=renderX)
        img12.grid(row=7, column=5)
        failureMessage11='Failure: Hall A Voltage, '+str(voltA)+'V'
        text.config(text=failureMessage11, bg='red', fg='black')
        passFail=0
        
    if voltB >=bHallVoltMin and voltB<=bHallVoltMax:
        img13=Label(gui, borderwidth=0, image=renderCheck)
        img13.grid(row=8, column=5)
    else:
        img13=Label(gui, borderwidth=0, image=renderX)
        img13.grid(row=8, column=5)
        failureMessage12='Failure: Hall B Voltage, '+str(voltB)+'V'
        text.config(text=failureMessage12, bg='red', fg='black')
        passFail=0

    if voltC >=cHallVoltMin and voltC<=cHallVoltMax:
        img14=Label(gui, borderwidth=0, image=renderCheck)
        img14.grid(row=9, column=5)
    else:
        img14=Label(gui, borderwidth=0, image=renderX)
        img14.grid(row=9, column=5)
        failureMessage13='Failure: Hall C Voltage, '+str(voltC)+'V'
        text.config(text=failureMessage13, bg='red', fg='black')
        passFail=0
     
    if Ha12 >=abHallMin and Ha12<=abHallMax:
        img15=Label(gui, borderwidth=0, image=renderCheck)
        img15.grid(row=10, column=5)
    else:
        img15=Label(gui, borderwidth=0, image=renderX)
        img15.grid(row=10, column=5)
        failureMessage14='Failure: Hall AB Alignment, '+str(Ha12)+'deg'
        text.config(text=failureMessage14, bg='red', fg='black')
        passFail=0
        
    if Ha13 >=acHallMin and Ha13<=acHallMax:
        img16=Label(gui, borderwidth=0, image=renderCheck)
        img16.grid(row=11, column=5)
    else:
        img16=Label(gui, borderwidth=0, image=renderX)
        img16.grid(row=11, column=5)
        failureMessage15='Failure: Hall AC Alignment, '+str(Ha13)+'deg'
        text.config(text=failureMessage15, bg='red', fg='black')
        passFail=0

    if Ha23 >=bcHallMin and Ha23<=bcHallMax:
        img17=Label(gui, borderwidth=0, image=renderCheck)
        img17.grid(row=12, column=5)
    else:
        img17=Label(gui, borderwidth=0, image=renderX)
        img17.grid(row=12, column=5)
        failureMessage16='Failure: Hall BC Alignment, '+str(Ha23)+'deg'
        text.config(text=failureMessage16, bg='red', fg='black')
        passFail=0

    print("HallA sensor Duty cycle : " +str(DCA)+"%")
    print("HallB sensor Duty cycle : " +str(DCB)+"%")
    print("HallC sensor Duty cycle : " +str(DCC)+"%")
    print("HallA sensor voltage level : "+str(aHallVolt[1])+" V")
    print("HallB sensor voltage level : "+str(bHallVolt[1])+" V")
    print("HallC sensor voltage level : "+str(cHallVolt[1])+" V")
    print('Hall A-B '+ str(Ph12) +' Deg')
    print('Hall A-C '+ str(Ph13) +' Deg')
    print('Hall B-C '+ str(Ph23) +' Deg')

def failureTableSheet():
    global failTable,failureMessage,failureMessage2,failureMessage3,failureMessage4,failureMessage5,failureMessage6,failureMessage7,failureMessage8,failureMessage9,failureMessage10,failureMessage11,failureMessage12,failureMessage13,failureMessage14,failureMessage15,failureMessage16,failureMessage17,failureMessage18
    failTable=[]
    row=[]
    row.append(failureMessage)
    failTable.append(row)
    row=[]
    row.append(failureMessage2)
    failTable.append(row)
    row=[]
    row.append(failureMessage3)
    failTable.append(row)
    row=[]
    row.append(failureMessage4)
    failTable.append(row)
    row=[]
    row.append(failureMessage5)
    failTable.append(row)
    row=[]
    row.append(failureMessage6)
    failTable.append(row)
    row=[]
    row.append(failureMessage7)
    failTable.append(row)
    row=[]
    row.append(failureMessage17)
    failTable.append(row)
    row=[]
    row.append(failureMessage18)
    failTable.append(row)
    row=[]
    row.append(failureMessage8)
    failTable.append(row)
    row=[]
    row.append(failureMessage9)
    failTable.append(row)
    row=[]
    row.append(failureMessage10)
    failTable.append(row)
    row=[]
    row.append(failureMessage11)
    failTable.append(row)
    row=[]
    row.append(failureMessage12)
    failTable.append(row)
    row=[]
    row.append(failureMessage13)
    failTable.append(row)
    row=[]
    row.append(failureMessage14)
    failTable.append(row)
    row=[]
    row.append(failureMessage15)
    failTable.append(row)
    row=[]
    row.append(failureMessage16)
    failTable.append(row)
    

def dataLog():
    global passFail,noHalls,failTable,barcode,folder,partNum,statorBarcode,speedMin,speedMax,aVrmsMin,aVrmsMax,bVrmsMin,bVrmsMax,cVrmsMin,cVrmsMax,abPhaseMin,abPhaseMax,acPhaseMin,acPhaseMax,bcPhaseMin,bcPhaseMax,aHallPhaseMin,aHallPhaseMax,bHallDutyMin,bHallPhaseMax,cHallDutyMin,cHallPhaseMax,aHallDutyMin,aHallDutyMax,bHallDutyMin,bHallDutyMax,cHallDutyMin,cHallDutyMax,aHallVoltMin,aHallVoltMax,bHallVoltMin,bHallVoltMax,cHallVoltMin,cHallVoltMax,abHallMin,abHallMax,acHallMin,acHallMax,bcHallMin,bcHallMax,rpm,rmsA,rmsB,rmsC,Ph12,Ph13,Ph23,BemfHallPh,BemfHallPhb,BemfHallPhc,DCA,DCB,DCC,voltA,voltB,voltC,Ha12,Ha13,Ha23
    table=[]
    row = []
    row.append(statorBarcode)
    table.append(row)
    row=[]
    row.append(rpm)
    row.append(speedMin)
    row.append(speedMax)
    table.append(row)
    try:
        row = []
        row.append(rmsA)
        row.append(aVrmsMin)
        row.append(aVrmsMax)
        table.append(row)
        row = []
        row.append(rmsB)
        row.append(bVrmsMin)
        row.append(bVrmsMax)
        table.append(row)
        row = []
        row.append(rmsC)
        row.append(cVrmsMin)
        row.append(cVrmsMax)
        table.append(row)
        row = []
        row.append(Ph12)
        row.append(abPhaseMin)
        row.append(abPhaseMax)
        table.append(row)
        row = []
        row.append(Ph13)
        row.append(acPhaseMin)
        row.append(acPhaseMax)
        table.append(row)
        row = []
        row.append(Ph23)
        row.append(bcPhaseMin)
        row.append(bcPhaseMax)
        table.append(row)
        if noHalls == 0:
            row = []
            row.append(BemfHallPh)
            row.append(aHallPhaseMin)
            row.append(aHallPhaseMax)
            table.append(row)
            row = []
            row.append(BemfHallPhb)
            row.append(bHallPhaseMin)
            row.append(bHallPhaseMax)
            table.append(row)
            row = []
            row.append(BemfHallPhc)
            row.append(cHallPhaseMin)
            row.append(cHallPhaseMax)
            table.append(row)
            row = []
            row.append(DCA)
            row.append(aHallDutyMin)
            row.append(aHallDutyMax)
            table.append(row)
            row = []
            row.append(DCB)
            row.append(bHallDutyMin)
            row.append(bHallDutyMax)
            table.append(row)
            row = []
            row.append(DCC)
            row.append(cHallDutyMin)
            row.append(cHallDutyMax)
            table.append(row)
            row = []
            row.append(voltA)
            row.append(aHallVoltMin)
            row.append(aHallVoltMax)
            table.append(row)
            row = []
            row.append(voltB)
            row.append(bHallVoltMin)
            row.append(bHallVoltMax)
            table.append(row)
            row = []
            row.append(voltC)
            row.append(cHallVoltMin)
            row.append(cHallVoltMax)
            table.append(row)
            row = []
            row.append(Ha12)
            row.append(abHallMin)
            row.append(abHallMax)
            table.append(row)
            row = []
            row.append(Ha13)
            row.append(acHallMin)
            row.append(acHallMax)
            table.append(row)
            row = []
            row.append(Ha23)
            row.append(bcHallMin)
            row.append(bcHallMax)
            table.append(row)
        if noHalls==0:
            resultsIndex=['StatorBarcode','Start RPM','Phase A Vrms','Phase B Vrms','Phase C Vrms','Phase AB Alignment','Phase AC Alignment','Phase BC Alignment','PhaseA HallA Alignment','PhaseB HallB Alignment','PhaseC HallC Alignment','Hall A Duty%','Hall B Duty%','Hall C Duty%','Hall A Volt','Hall B Volt','Hall C Volt','Hall AB Alignment','Hall AC Alignment','Hall BC Alignment']
        else:
            resultsIndex=['StatorBarcode','Start RPM','Phase A Vrms','Phase B Vrms','Phase C Vrms','Phase AB Alignment','Phase AC Alignment','Phase BC Alignment']
    except:
        resultsIndex=['Start RPM'] # if exception occurs RPM has failed, only list those variables
    resultsCols = ['Measured Values', 'Minimum','Maximum'] # column labels
    failureTableSheet()
    dfResults = pd.DataFrame(table, columns=resultsCols, index=resultsIndex)# place table into a pandas dataframe object
    failureCols = ['Failure Message'] # column labels
    failureIndex=['Message1','Message2','Message3','Message4','Message5','Message6','Message7','Message8','Message9','Message10','Message11','Message12','Message13','Message14','Message15','Message16','Message17','Message18']
    dfFailures = pd.DataFrame(failTable, columns=failureCols, index=failureIndex) #index=resultsIndex) # place table into a pandas dataframe object
    timestr = time.strftime("(%Y%m%d-%H%M%S)") # add date time stamp and save location
    if passFail==1:
        file = barcode+timestr # file name with date time stamp
    elif passFail==0:
        file = 'Failure-'+barcode+timestr
    extension = '.xlsx' # file extension
    if os.path.isdir(folder+'/'+partNum)==False:
        os.mkdir(folder+'/'+partNum)
    fullPath = pathlib.Path(os.path.join(folder+'/'+partNum+'/'+file+extension)) # the full file path
    writer = pd.ExcelWriter(path=fullPath, engine='xlsxwriter',engine_kwargs={'options': {'strings_to_formulas': False}})
    dfResults.to_excel(writer, sheet_name='Results')
    dfFailures.to_excel(writer, sheet_name='Failure Messages')
    writer.save()
    
def savePlot():
    global barcode,saveIndex,passFail,screenshotsSavePath,partNum
    if saveIndex==1: # 1 = main tests are finished

        screenshotsSavePath = configTemplate.screenshotsSavePath # screen shot save path
        screenshotsSavePath = screenshotsSavePath+'/'+partNum
        if os.path.isdir(screenshotsSavePath):
            print('Screenshot directory found')
        else:
            os.mkdir(screenshotsSavePath)
        timestr = time.strftime("(%Y%m%d-%H%M%S)")
        file = barcode+timestr
        if passFail==0:
            file='Failure-'+file
        fullPath = pathlib.Path(os.path.join(screenshotsSavePath, file+'.rtf'))
        fileNames = ['HallPhaseA_To_MotorPhaseA_Alignment.png','HallPhaseB_To_MotorPhaseB_Alignment.png','HallPhaseC_To_MotorPhaseC_Alignment.png','MotorPhases.png','HallPhases.png']
        i=0
        doc = docx.Document() # Create an instance of a word document
        doc.add_heading(barcode, 1) # Add a Title to the document
        for files in fileNames:
            doc.add_heading(fileNames[i].strip('.png'), 3) # Image in its native size
            doc.add_picture(files, width=Inches(6), height=Inches(3.5))
            #document.add_page_break()
            i=i+1
        doc.save(fullPath) # Now save the document to a location
        os.startfile(str(fullPath))
        saveIndex=0
    gui.after(500,savePlot)

def openResults():
    filename = filedialog.askopenfilename(initialdir=folder, title="BackEMF Test Results",filetypes=(("all files", "*.*"),("xlsx files", "*.xlsx"),("image files", "*.png"),("image files", "*.bmp")))
    ext = os.path.splitext(filename)
    ext=ext[1]
    if len(filename) >=3 and ext == '.xlsx' :
        path = pathlib.Path(filename)
        os.startfile(path)

def openReports():
    filename = filedialog.askopenfilename(initialdir=screenshotsSavePath, title="BackEMF Test Reports",filetypes=(("all files", "*.*"),("rich text files", "*.rtf"),("image files", "*.png"),("image files", "*.bmp")))
    ext = os.path.splitext(filename)
    ext=ext[1]
    if len(filename) >=3 and ext == '.rtf' :
        path = pathlib.Path(filename)
        os.startfile(path)
        
def openConfig():
    global configFile
    configFile = filedialog.askopenfilename(initialdir=configFolder, title="BackEMF Config Files",filetypes=(("txt files", "*.txt"),("all files", "*.*")))
    configFile = '"'+configFile+'"'
    if len(configFile) >=3:
        path = str('start '+ "notepad " + configFile)
        os.system(path)
       
def startSpeedWindow():
    startSpeed = Toplevel(gui)
    startSpeed.geometry("400x150")
    startSpeed.title("Start Drill Press")
    startSpeed.wm_iconbitmap('icon.ico')
    startSpeed.config(bg='blue')
    startSpeedLabel=tk.Label(startSpeed, text="Start Drill Press", height=3, width=22, font='12', fg='white', bg='blue')
    startSpeedLabel.grid(row=0, column=0, columnspan=6, ipadx=100)
    startSpeedLabel.update()
    def cancelSpeedWindow():
        resetGui()
        startSpeed.destroy()
    def getSpeedScope1(): # start test
        global failureMessage,status,rpm,passFail,Scope1,speedMin,speedMax,rpm,last,cancelIndex#,startSpeedWindow
        Scope=Scope1
        rpm=0.0
        while rpm <= 10 or rpm >= 300: # exit if motor is moving
            text.config(text='Start Drill Press', bg='orange', fg='black')
            rpm=round(SDS1104XE.measureFreq(Scope, 'C2')/26*60)
            if cancelIndex == 0:
                status='Cancelled, Please Enter Barcode'
                text.config(text=status, bg='#3399ff', fg='black')
                passFail=0
                break
        if passFail==1:
            rpm=round(SDS1104XE.measureFreq(Scope, 'C2')/26*60)
            print("Rotor speed : " + str(rpm) + " RPM") # rpm=freq/(rotorMagnetCount/2)*60sec
            if rpm >= speedMin and rpm <= speedMax: # if final speed is in tolerance pass
                text.config(text='Start Speed Pass: '+str(rpm)+ "RPM", bg='orange', fg='black')
            else:
                failureMessage='Failure: Start Speed '+str(rpm)+ "RPM is too low or too high"
                text.config(text=failureMessage, bg='red', fg='black')
                passFail=0
        startSpeed.destroy()
        return
    cancel_button1 = tk.Button(startSpeed, font='12',width=20, text="Cancel", command=cancelSpeedWindow)
    cancel_button1.grid(row=2, column=0, columnspan=6)
    cancel_button1.update()
    getSpeedScope1()
    return

def statorBarcodeWindow():
    global statorBarcode,statorBarcode_var,skipButton
    statorBarcodeWindow= Toplevel(gui)
    statorBarcodeWindow.geometry("400x200")
    statorBarcodeWindow.title("Enter Stator Barcode")
    statorBarcodeWindow.wm_iconbitmap('icon.ico')
    statorBarcodeWindow.config(bg='#3399ff')
    #Create a Label in New window
    statorLabel=tk.Label(statorBarcodeWindow, text="Enter Stator Barcode", height=3, width=22, font='12', fg='black', bg='#3399ff')
    statorLabel.grid(row=1, column=0, columnspan=6, ipadx=100, ipady=5)
    statorLabel.update()
    def skipBarcode():
        global statorBarcode,stator
        statorBarcode=stator+'-xxxx'
        statorBarcodeWindow.destroy()
        return statorBarcode
    def destroyWindow(event):
        global statorBarcode
        statorBarcode=statorBarcode_var.get()
        statorBarcodeWindow.destroy()
        return statorBarcode
    def cancelWindow():
        resetGui()
        statorBarcodeWindow.destroy()
    statorBarcode_var = tk.StringVar() 
    expression_field1=tk.Entry(statorBarcodeWindow, font=(12), relief="groove", textvariable=statorBarcode_var, borderwidth=10)
    expression_field1.bind('<Return>',destroyWindow)
    expression_field1.grid(row=0, column=0, columnspan=6, ipadx=100)
    expression_field1.icursor('end')
    expression_field1.focus_set()
    expression_field1.update()
    statorBarcode_var.set('')
    cancel_button = tk.Button(statorBarcodeWindow,width=20,font='12',text="Cancel",command=cancelWindow)
    cancel_button.grid(row=2, column=0, columnspan=6)
    cancel_button.update()
    if skipButton == '1':
        skip_button = tk.Button(statorBarcodeWindow,width=20,font='12',text="Skip",command=skipBarcode)
        skip_button.grid(row=3, column=0, columnspan=6)
        skip_button.update()
    return statorBarcode

def getStatorInfo(): # event is used for gui entry .bind <return>
    global statorBarcode,partNum,passFail,status, cancelIndex,stator,statorBarcode_var
    try:
        statorBarcodeWindow()
        cancelIndex = 1 # for cancel 
        while statorBarcode == '':
            if cancelIndex == 0: # status gets reset by cancel_button
                break
        statorBarcode=statorBarcode.upper() # converts barcode to all upercase
        print("Stator Barcode="+statorBarcode)
        if len(statorBarcode.split('-')) == 4 and cancelIndex == 1: # check if the barcode is the correct length/format
            first, second, third, sn = statorBarcode.split('-') # put barcode into list of strings
            statorPartNum=first+'-'+second+'-'+third
            if statorPartNum==stator:
                text.config(fg="black",bg='Yellow',width=25,font=12, text='Testing: '+partNum)
                barcode_var.set('')
            else:
                raise ValueError # if wrong format, ask to scan again
        elif cancelIndex == 0:
            status='Cancelled, Please Enter Barcode'
            text.config(text=status, bg='#3399ff', fg='black')
            passFail=0
            
        else:
            raise ValueError # if wrong format, ask to scan again
            
    except ValueError:
        print('ERROR: Invalid Stator PN! Please Start Over.')
        status='Error: Invalid Stator PN! Please Start Over.'
        text.config(fg="black",bg='red',width=25,font=12, text=status)
        barcode_var.set('')
        passFail=0
    return passFail
        
def about():
    window = tk.Tk()
    window.title("About")
    window.rowconfigure(0, minsize=400, weight=1)      #about window row#, size, wieght
    window.wm_iconbitmap('icon.ico')
    window.geometry('1200x600')
    window.config(bg='#3399ff')
    message ='''
    --About the program--
    Back EMF Tester GUI
    Author: Doug Gammill
    Date: 7/26/2022
    Copyright (c) 2022
    Linear Labs, Inc.
    Tests Motor Back EMF; For final assembly or motor rework.
    Phase Alignment, Phase Voltage, Hall Alignment, Hall Duty Cycle, Hall to Phase Alignment.
    General Test Procedure:
    1. Turn on Scope1 and Scope2.
    2. Ensure all USB cables for each Scope and USB relays are connected to the PC
    3. Ensure Linear Labs Back EMF computer application (gui) has indicated USB coms are good with each device.
    4. Connect the desired part to the test fixture. Make sure to use the correct cable adapters for the part(s).
    5. Scan the barcode located on the motor using the barcode scanner.
    6. Next a pop up window will appear, scan stator SN here or click "skip"M10-000135-002-0494
    .
    7. Turn on drill press to spin motor. 
    8. Confirm that there are no failure(s) for the part by reading the label outputs.
    9. Check the excel output file if desired by clicking the 'Results' in the options menu.
    '''
    text_box = Text(window,height=23,width=95, font=("Adobe",16))
    text_box.pack(expand=True)
    text_box.insert('end', message)
    text_box.config(state='disabled')
    window.mainloop()
    
def configPaths():
    os.startfile('folderConfig.txt')
    
def failure():
    global failureMessage,failureMessage2,failureMessage3,failureMessage4,failureMessage5,failureMessage6,failureMessage7,failureMessage8,failureMessage9,failureMessage10,failureMessage11,failureMessage12,failureMessage13,failureMessage14,failureMessage15,failureMessage16,failureMessage17,failureMessage18
    #sleep(3)
    fail = tk.Tk()
    fail.title("Failure!!!")
    fail.rowconfigure(0, minsize=400, weight=1)      #about window row#, size, wieght
    fail.wm_iconbitmap('icon.ico')
    fail.geometry('600x500')
    fail.config(bg='red')
    ####if there is a message insert newline at end
    if failureMessage != '':
        failureMessage=failureMessage+'\n'
    if failureMessage2 != '':
        failureMessage2=failureMessage2+'\n'       
    if failureMessage3 != '':
        failureMessage3=failureMessage3+'\n'       
    if failureMessage4 != '':
        failureMessage4=failureMessage4+'\n'        
    if failureMessage5 != '':
        failureMessage5=failureMessage5+'\n'       
    if failureMessage6 != '':
        failureMessage6=failureMessage6+'\n'       
    if failureMessage7 != '':
        failureMessage7=failureMessage7+'\n'       
    if failureMessage8 != '':
        failureMessage8=failureMessage8+'\n'      
    if failureMessage9 != '':
        failureMessage9=failureMessage9+'\n'      
    if failureMessage10 != '':
        failureMessage10=failureMessage10+'\n'       
    if failureMessage11 != '':
        failureMessage11=failureMessage11+'\n'       
    if failureMessage12 != '':
        failureMessage12=failureMessage12+'\n'      
    if failureMessage13 != '':
        failureMessage13=failureMessage13+'\n'       
    if failureMessage14 != '':
        failureMessage14=failureMessage14+'\n'
    if failureMessage15 != '':
        failureMessage15=failureMessage15+'\n'
    if failureMessage16 != '':
        failureMessage16=failureMessage16+'\n'
    if failureMessage17 != '':
        failureMessage17=failureMessage17+'\n'
        
    message = failureMessage+failureMessage2+failureMessage3+failureMessage4+failureMessage5+failureMessage6+failureMessage7+failureMessage8+failureMessage9+failureMessage10+failureMessage11+failureMessage12+failureMessage13+failureMessage14+failureMessage15+failureMessage16+failureMessage17+failureMessage18

    text_box = Text(fail,height=18,width=46, font=("Adobe",16))
    text_box.pack(expand=True)
    text_box.insert('end', message)
    text_box.config(state='disabled')

    def reset():
        text.config(fg="black",bg='#3399ff',width=25,font=12, text='Please Enter Barcode Bellow')
        fail.destroy()
        
    btn_ok = tk.Button(fail, font=("Adobe",16),width=10, text="OK", command=reset)
    btn_ok.pack()
    fail.mainloop()  
    
def getBarcodeInfo(): # event is used for gui entry .bind <return>
    global barcode, partNum, sn, passFail, status
    barcode=barcode_var.get()
    barcode=barcode.upper() # converts barcode to all upercase
    usbComs() # check coms real quick
    testIndex=1 # determines if testing or not 1 = testing
    try:
        print(barcode)                                                
        if len(barcode.split('-')) == 4: # check if the barcode is the correct length/format
            first, second, third, sn = barcode.split('-') # put barcode into list of strings
            partNum=first+'-'+second+'-'+third
            text.config(fg="white",bg='Purple',width=25,font=12, text='Please Scan Installed Stator Barcode')
            barcode_var.set('')
        else:
            raise ValueError # if wrong format, ask to scan again
    except ValueError:
        print('ERROR: Invalid barcode! Please scan again.')
        status='Error: Invalid Barcode! Please Scan Again.'
        text.config(fg="black",bg='red',width=25,font=12, text=status)
        barcode_var.set('')
        passFail=0
        return passFail
# Driver code 
#if __name__ == "__main__":
# create a GUI window
gui = Tk()
for i in range(12):  ###rows = 12
    gui.columnconfigure(i, weight=1, minsize=100)
    gui.rowconfigure(i, weight=1, minsize=60)

    for j in range(0, 6): ###column = 7
        frame = tk.Frame(
            master=gui,
            relief=tk.RAISED,
            borderwidth=1
        )

# set the background colour of GUI window
gui.configure(background="Black")

# set the title of GUI window
gui.title("Back EMF Tester")
gui.wm_iconbitmap('icon.ico')

text_var = tk.StringVar()
text=tk.Label(gui, font=12, relief="ridge", text='Setting Up USB, Please Wait', borderwidth=10)
text.grid(row=1, columnspan=6, ipadx=100)
text.update()

# set the configuration of GUI window
gui.geometry("600x800") # width by length
gui.resizable(width=True, height=True)

def resetGui():
    global statorBarcode,cancelIndex,img2,img3,img4,img5,img6,img7,img8,img9,img10,img11,img12,img13,img14,img15,img16,img17,img18,img19,failureMessage,failureMessage2,failureMessage3,failureMessage4,failureMessage5,failureMessage6,failureMessage7,failureMessage8,failureMessage9,failureMessage10,failureMessage11,failureMessage12,failureMessage13,failureMessage14,failureMessage15,failureMessage16,failureMessage17,failureMessage18,status,speedMin,speedMax,aVrmsMin,aVrmsMax,bVrmsMin,bVrmsMax,cVrmsMin,cVrmsMax,abPhaseMin,abPhaseMax,acPhaseMin,acPhaseMax,bcPhaseMin,bcPhaseMax,aHallPhaseMin,aHallPhaseMax,aHallDutyMin,aHallDutyMax,bHallDutyMin,bHallDutyMax,cHallDutyMin,cHallDutyMax,aHallVoltMin,aHallVoltMax,bHallVoltMin,bHallVoltMax,cHallVoltMin,cHallVoltMax,abHallMin,abHallMax,acHallMin,acHallMax,bcHallMin,bcHallMax
    img2.destroy()
    img3.destroy()
    img4.destroy()
    img5.destroy()
    img6.destroy()
    img7.destroy()
    img8.destroy()
    img9.destroy()
    img10.destroy()
    img11.destroy()
    img12.destroy()
    img13.destroy()
    img14.destroy()
    img15.destroy()
    img16.destroy()
    img17.destroy()
    img18.destroy()
    img19.destroy()
    speedMin = ''
    speedMax = ''
    aVrmsMin = ''
    aVrmsMax = ''
    bVrmsMin = ''
    bVrmsMax = ''
    cVrmsMin = ''
    cVrmsMax = ''
    abPhaseMin = ''
    abPhaseMax = ''
    acPhaseMin = ''
    acPhaseMax = ''
    bcPhaseMin = ''
    bcPhaseMax = ''
    aHallPhaseMin = ''
    aHallPhaseMax = ''
    aHallDutyMin = ''
    aHallDutyMax = ''
    bHallDutyMin = ''
    bHallDutyMax = ''
    cHallDutyMin = ''
    cHallDutyMax = ''
    aHallVoltMin = ''
    aHallVoltMax = ''
    bHallVoltMin = ''
    bHallVoltMax = ''
    cHallVoltMin = ''
    cHallVoltMax = ''
    abHallMin = ''
    abHallMax = ''
    acHallMin = ''
    acHallMax = ''
    bcHallMin = ''
    bcHallMax = ''
    failureMessage=''
    failureMessage2=''
    failureMessage3=''
    failureMessage4=''
    failureMessage5=''
    failureMessage6=''
    failureMessage7=''
    failureMessage8=''
    failureMessage9=''
    failureMessage10=''
    failureMessage11=''
    failureMessage12=''
    failureMessage13=''
    failureMessage14=''
    failureMessage15=''
    failureMessage16=''
    failureMessage17=''
    failureMessage18=''
    status=''
    statorBarcode=''
    cancelIndex = 0

def checkUSB():
    global comOK,testIndex,passFail1, passFail2, passFailArduino
    if comOK == 1 and testIndex == 0: # comOK=1 means at some point USB was good and index=0 means not currently testing
        text.config(fg="black",bg='orange',text='Checking USB')
        try:
            SDS1104XE.checkScope1()
            passFail1=SDS1104XE.passFail1
            if passFail1 == 1:
                    scope1usb.config(fg="black",bg='red',text='Connect USB')
                    comOK = 0  # fail
                    
            SDS1104XE.checkScope2()
            passFail2=SDS1104XE.passFail2
            if passFail2 == 1:
                    scope2usb.config(fg="black",bg='red',text='Connect USB')
                    comOK = 0  # fail
                    
            arduino.checkArduino()
            passFailArduino=arduino.passFailArduino
            if passFailArduino == 1:
                    relay1usb.config(fg="black",bg='red',text='Connect USB')
                    comOK = 0  # fail
            if passFail1 == 0 and passFail2 == 0 and passFailArduino == 0:
                text.config(fg="black",bg='#3399ff',text='Enter Barcode')
            else:
                text.config(bg='red', fg='black', text='Connect USB, Scope Restart may be required')
     
        except:
            comOK=0
            print("something went wrong")
            text.config(fg="black",bg='red',text='Connect USB, Scope Restart may be required')

def connectUSB():
    global Scope1, Scope2, passFail1,passFail2,passFailArduino,comOK, testIndex
    while comOK == 0 and testIndex == 0: # comOK=1 means at some point USB was good and index=0 means not currently testing

        SDS1104XE.connectOScope1()
        passFail1=SDS1104XE.passFail1
        if passFail1==1: # 1 = failure
            scope1usb.config(bg='red', fg='black', text='Connect USB')
        else: # pass
            Scope1=SDS1104XE.my_instrument1
            scope1usb.config(bg='purple',fg='white', text='Scope 1')
            
        SDS1104XE.connectOScope2()
        passFail2=SDS1104XE.passFail2
        if passFail2==1: # 1 = failure
            scope2usb.config(bg='red', fg='black', text='Connect USB')
        else:
            Scope2=SDS1104XE.my_instrument2
            scope2usb.config(bg='purple',fg='white', text='Scope 2')
            
        arduino.connectArduino()
        passFailArduino = arduino.passFailArduino
        if passFailArduino==1: # 1 = failure
            relay1usb.config(bg='red',fg='black', text='Connect USB')
        else:
            relay1usb.config(bg='purple',fg='white', text='Relays')
            
        if passFail1 == 0 and passFail2 == 0 and passFailArduino == 0:
            comOK = 1 # 1 = passing
        else:
            text.config(bg='red', fg='black', text='Connect USB, Scope Restart may be required')

def usbComs():
    checkUSB()
    connectUSB()
    checkUSB()
    return

def usbThread():
    t1=Thread(target=usbComs)
    t1.start()
    
def threading(event): # threading used to execute main functions and allow GUI main loop to not freeze
    t1=Thread(target=mainFunctions)
    t1.start()   

def mainFunctions(): # main multi threaded functions
    global passFail,passFail1,passFail2,status,skipFail,noHalls,saveIndex,testIndex
    passFail=1
    resetGui()
    getBarcodeInfo()
    if passFail==1:
        readConfig() # read variables from config file
    if passFail==1:
        getStatorInfo()
    if passFail==1:
        setupScopes()
        #sleep(1)
    if passFail==1:
        startSpeedWindow()
    #if passFail==1:
        #autoSetup() # same as auto button on scope
    if passFail==1 and skipFail==0: # will perform full test even if failure happens
        getValuesScope1()
        if noHalls == 0:
            getValuesScope2()
    if passFail==1 and skipFail==1: # will skip scope2 test if set to 1 in configFile
        getValuesScope1()
        if passFail==1 and noHalls==0:
            getValuesScope2()
    if passFail==1:
        saveIndex=1
        dataLog()
        text.config(text='Test Passed Please Enter Barcode Bellow', bg = '#3399ff')
    if passFail==0 and status=='':
        saveIndex=1  # save plot index
        testIndex=0  # determines if testing or not 1 = testing
        dataLog()
        failure()
    return

barcode_var = tk.StringVar() 
expression_field = tk.Entry(gui, font=(12), relief="groove", textvariable=barcode_var, borderwidth=10)
expression_field.grid(row=2, column=0, columnspan=6, ipadx=100)
expression_field.bind('<Return>',threading)
expression_field.icursor('end')
expression_field.focus_set() # insert cursor
expression_field.update()

menubar = Menu(gui, background='#ff8000', foreground='black', activebackground='white', activeforeground='black')
help = Menu(menubar, tearoff=0)  
help.add_command(label="About", command=about)
help.add_command(label="Results", command=openResults)
help.add_command(label="Reports", command=openReports)
help.add_command(label="Configs", command=openConfig)
help.add_command(label="Paths", command=configPaths)
#help.add_command(label="Check USB", command=usbThread)# only use during debugging, pressing multiple times leads to multiple threads.
menubar.add_cascade(label="Options", menu=help)
gui.config(menu=menubar)

scope1usb=tk.Label(gui, text="Scope1", height=3, width=16, font='12', fg='black', bg='white', relief="groove", borderwidth=2)
scope1usb.grid(row=0, column=0)
scope1usb.update()
scope2usb=tk.Label(gui, text="Scope2", height=3, width=16, font='12', fg='black', bg='white', relief="groove", borderwidth=2)
scope2usb.grid(row=0, column=1)
scope2usb.update()
relay1usb=tk.Label(gui, text="Relays", height=3, width=16, font='12', fg='black', bg='white', relief="groove", borderwidth=2)
relay1usb.grid(row=0, column=2)
relay1usb.update()

test1=tk.Label(gui, text="Scope1 Test:", height=3, width=20, font="Verdana 12 underline", fg='white', bg='blue', relief="groove")
test1.grid(row=3, column=0, columnspan=2)
test1.update()
phase1=tk.Label(gui, text="Phases", height=3, width=16, font="Verdana 12 underline", fg='white', bg='blue', relief="groove")
phase1.grid(row=3, column=2)
phase1.update()
test2=tk.Label(gui, text="Scope2 Test:", height=3, width=20, font="Verdana 12 underline", fg='white', bg='blue', relief="groove")
test2.grid(row=3, column=3, columnspan=2)
test2.update()
hall2=tk.Label(gui, text="Halls", height=3, width=16, font="Verdana 12 underline", fg='white', bg='blue', relief="groove")
hall2.grid(row=3, column=5)
hall2.update()

VA=tk.Label(gui, text="Phase A Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
VA.grid(row=4, column=0, columnspan=2)
VA.update()
VB = tk.Label(gui, text="Phase B Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
VB.grid(row=5, column=0, columnspan=2)
VB.update()
VC = tk.Label(gui, text="Phase C Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
VC.grid(row=6, column=0, columnspan=2)
VC.update()
PAB = tk.Label(gui, text="Phase AB  Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
PAB.grid(row=7, column=0, columnspan=2)
PAB.update()
PAC = tk.Label(gui, text="Phase AC Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
PAC.grid(row=8, column=0, columnspan=2)
PAC.update()
PBC = tk.Label(gui, text="Phase BC Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
PBC.grid(row=9, column=0, columnspan=2)
PBC.update()
HPA = tk.Label(gui, text="HallA PhaseA Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HPA.grid(row=10, column=0, columnspan=2)
HPA.update()
HPB = tk.Label(gui, text="HallB PhaseB Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HPB.grid(row=11, column=0, columnspan=2)
HPB.update()
HPC = tk.Label(gui, text="HallC PhaseC Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HPC.grid(row=12, column=0, columnspan=2)
HPC.update()

DutyA = tk.Label(gui, text="Hall A Duty Cycle", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
DutyA.grid(row=4, column=3, columnspan=2)
DutyA.update()
DutyB = tk.Label(gui, text="Hall B Duty Cycle", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
DutyB.grid(row=5, column=3, columnspan=2)
DutyB.update()
DutyC=tk.Label(gui, text="Hall C Duty Cycle", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
DutyC.grid(row=6, column=3, columnspan=2)
DutyC.update()
HVA=tk.Label(gui, text="Hall A Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HVA.grid(row=7, column=3, columnspan=2)
HVA.update()
HVB = tk.Label(gui, text="Hall B Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HVB.grid(row=8, column=3, columnspan=2)
HVB.update()
HVC = tk.Label(gui, text="Hall C Voltage", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HVC.grid(row=9, column=3, columnspan=2)
HVC.update()
HAB = tk.Label(gui, text="Hall AB Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HAB.grid(row=10, column=3, columnspan=2)
HAB.update()
HAC = tk.Label(gui, text="Hall AC Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HAC.grid(row=11, column=3, columnspan=2)
HAC.update()
HBC = tk.Label(gui, text="Hall BC Alignment", height=3, width=22, font='12', fg='black', bg='#3399ff', relief="groove")
HBC.grid(row=12, column=3, columnspan=2)
HBC.update()

####IMG####
load= Image.open("logo4.ppm")
render = ImageTk.PhotoImage(load)
img = Label(gui, borderwidth=0, image=render)
img.grid(row=0, column=3, columnspan=3)
img.update()

load= Image.open("motorIcon.png")
render2 = ImageTk.PhotoImage(load)

img2=Label(gui, borderwidth=0, image=render2)
img2.grid(row=4, column=2)
img3=Label(gui, borderwidth=0, image=render2)
img3.grid(row=5, column=2)
img4=Label(gui, borderwidth=0, image=render2)
img4.grid(row=6, column=2)
img5=Label(gui, borderwidth=0, image=render2)
img5.grid(row=7, column=2)
img6=Label(gui, borderwidth=0, image=render2)
img6.grid(row=8, column=2)
img7=Label(gui, borderwidth=0, image=render2)
img7.grid(row=9, column=2)
img8=Label(gui, borderwidth=0, image=render2)
img8.grid(row=10, column=2)

img18=Label(gui, borderwidth=0, image=render2)
img18.grid(row=11, column=2)
img19=Label(gui, borderwidth=0, image=render2)
img19.grid(row=12, column=2)

img9=Label(gui, borderwidth=0, image=render2)
img9.grid(row=4, column=5)
img10=Label(gui, borderwidth=0, image=render2)
img10.grid(row=5, column=5)
img11=Label(gui, borderwidth=0, image=render2)
img11.grid(row=6, column=5)
img12=Label(gui, borderwidth=0, image=render2)
img12.grid(row=7, column=5)
img13=Label(gui, borderwidth=0, image=render2)
img13.grid(row=8, column=5)
img14=Label(gui, borderwidth=0, image=render2)
img14.grid(row=9, column=5)
img15=Label(gui, borderwidth=0, image=render2)
img15.grid(row=10, column=5)
img16=Label(gui, borderwidth=0, image=render2)
img16.grid(row=11, column=5)
img17=Label(gui, borderwidth=0, image=render2)
img17.grid(row=12, column=5)

load= Image.open("Red-X.ppm")
renderX = ImageTk.PhotoImage(load)
load= Image.open("Blue-Check1.ppm")
renderCheck = ImageTk.PhotoImage(load)
    
usbThread()
gui.after(100, savePlot)
gui.mainloop() # start the GUI
